from docx import Document
import random
import string
import pandas as pd
import os
import re
import shutil

directory_path = "M:\CDB\Analyst\Rhys\Python\CustodianExtract\custodian_extraction\input\BNP Paribase Singapore"
output_directory_path = "M:\CDB\Analyst\Rhys\Python\CustodianExtract\custodian_extraction\output\BNP Paribas Singapore"

company_names = ["GALAPAGOS NV", "TECK RESOURCES LTD", "IMPLENIA AG", "BANCO SANTANDER REG. SHS", "BANCO SANTANDER ADR REG. SHS", "KONE OYJ -B- REG. SHS", "DZ PRIVATBANK S.A.", "ARCELORMITTAL SA", "ASML HOLDING NV", "TAIWAN SEMICONDUCTOR MANUFACTURING CO., LTD.", "ARCELORMITTAL SA ADR", "ASML HOLDING NV ADR"]
random_names = ["AMAZON.COM INC", "FACEBOOK INC.", "JOHNSON & JOHNSON","ALPHABET INC.", "APPLE INC.", "GOOGLE INC.", "MICROSOFT CORP.", "AMERICAN AIRLINES", "ZARA", "MCDONALDS", "BRITISH AIRWAYS", "MERCEDES"]

df = pd.DataFrame({
    'Company': company_names,
    'Random_Name': random_names
})
print(df)

def extract_isin_codes(filename):
    # Define a regular expression pattern to match ISIN codes
    pattern = re.compile(r'\b[A-Za-z]{2}[0-9A-Za-z]{10,16}\b')

    match = pattern.search(filename)
    if match:
        isin = match.group()
        print(f"Found ISIN in the defined format: {isin}")
        return isin
    else:
        print("No ISIN found in the defined format.")
        return None

def find_and_replace(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_text in cell.text:
                    cell.text = cell.text.replace(old_text, new_text)

def replace_company_names(doc, company_names_df):
    for company_name in df['Company']:
        for paragraph in doc.paragraphs:
            if company_name in paragraph.text:
                # Find the corresponding row in the DataFrame
                match_row = df[df['Company'] == company_name]
                if not match_row.empty:
                    dummy_company = match_row['Random_Name'].values[0]

                    # Replace the company name in the paragraph
                    paragraph.text = paragraph.text.replace(company_name, dummy_company)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if company_name in cell.text:
                        # Find the corresponding row in the DataFrame
                        match_row = company_names_df[company_names_df['Company'] == company_name]
                        if not match_row.empty:
                            dummy_company = match_row['Random_Name'].values[0]

def generate_random_sequence(isin_format):
    letters = [random.choice(string.ascii_uppercase) for _ in range(2)]
    digits = [random.choice(string.digits) for _ in range(10)]
    return ''.join(letters + digits)

def wordReplaceWord(directory_path, output_directory_path, df):
    print("Files in the directory:")

    for filename in os.listdir(directory_path):
        full_path = os.path.join(directory_path, filename)
        print(full_path)
        doc = Document(full_path)
        old_word = extract_isin_codes(filename)
        print(old_word)
        new_word = generate_random_sequence(old_word)
        replace_company_names(doc, df)
        find_and_replace(doc, old_word, new_word)
        new_filename = f"ENCRYPTED_{new_word}.docx"
        doc.save("M:\\CDB\\Analyst\\Rhys\\Python\\CustodianExtract\\custodian_extraction\\input\\BNP Paribase Singapore\\" + new_filename)
        new_full_path = os.path.join(output_directory_path, os.path.splitext(filename)[0] + ".xlsx")
        output_new = os.path.join(output_directory_path, f"ENCRYPTED_{new_word}.xlsx")

        if os.path.exists(new_full_path):
            shutil.copy2(new_full_path, output_new)
            print(f"File copied to: {output_new}")
        else:
            print(f"File not found at: {new_full_path}")
        print(output_new)


wordReplaceWord(directory_path, output_directory_path, df)

